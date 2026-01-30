"""
Microsoft Word document parser.
"""
from pathlib import Path
from typing import List

from docx import Document

from src.core.logging import get_logger
from src.core.exceptions import DocumentParseError
from src.parsers.base import Parser


logger = get_logger(__name__)


class DocxParser(Parser):
    """
    Parser for Microsoft Word documents (.docx).
    
    This parser handles .docx files, extracting text content while
    attempting to preserve document structure like paragraphs, headings,
    and list items.
    
    Supported extensions: .docx
    
    Example:
        >>> parser = DocxParser()
        >>> text = parser.parse("document.docx")
        >>> print(text)
    """
    
    supported_extensions = ["docx"]
    
    def __init__(self, preserve_structure: bool = True) -> None:
        """
        Initialize DOCX parser.
        
        Args:
            preserve_structure: Whether to preserve document structure
                            (headings, lists, etc.)
        """
        self.preserve_structure = preserve_structure
    
    def parse(self, filepath: Path | str) -> str:
        """
        Parse a DOCX file and extract its text content.
        
        Args:
            filepath: Path to DOCX file
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentParseError: If file cannot be read or parsed
            
        Example:
            >>> parser = DocxParser()
            >>> text = parser.parse("document.docx")
        """
        path = Path(filepath)
        
        # Validate file
        self.validate_file(path)
        
        logger.info("Parsing DOCX file", filepath=str(path))
        
        try:
            # Open document
            doc = Document(path)
            
            # Extract text based on settings
            if self.preserve_structure:
                full_text = self._extract_with_structure(doc)
            else:
                full_text = self._extract_simple(doc)
            
            if not full_text:
                logger.warning("No text extracted from DOCX", filepath=str(path))
                return ""
            
            # Clean up text
            full_text = self._clean_text(full_text)
            
            logger.info(
                "Successfully parsed DOCX file",
                filepath=str(path),
                char_count=len(full_text),
                paragraph_count=len(doc.paragraphs)
            )
            
            return full_text
            
        except Exception as e:
            logger.error("Failed to parse DOCX file", filepath=str(path), error=str(e))
            raise DocumentParseError(
                f"Failed to parse DOCX file: {str(e)}",
                details={"filepath": str(filepath), "error": str(e)}
            )
    
    def _extract_with_structure(self, doc: Document) -> str:
        """
        Extract text while preserving document structure.
        
        This method preserves headings, list items, and paragraphs
        with appropriate formatting markers.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            Text with structure preserved
        """
        paragraphs = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            if not text:
                continue
            
            # Check if this is a heading
            if paragraph.style.name.startswith('Heading'):
                level = paragraph.style.name.replace('Heading ', '')
                # Add heading marker
                heading_marker = '#' * int(level) if level.isdigit() else '#'
                paragraphs.append(f"{heading_marker} {text}")
            
            # Check if this is a list item
            elif paragraph.style.name.startswith('List'):
                # Determine bullet or numbered list
                if 'Bullet' in paragraph.style.name:
                    paragraphs.append(f"• {text}")
                else:
                    paragraphs.append(f"1. {text}")
            
            # Regular paragraph
            else:
                paragraphs.append(text)
        
        # Add tables if present
        if doc.tables:
            paragraphs.append("\n[Tables]")
            for table in doc.tables:
                table_text = self._extract_table(table)
                if table_text:
                    paragraphs.append(table_text)
        
        return '\n\n'.join(paragraphs)
    
    def _extract_simple(self, doc: Document) -> str:
        """
        Extract text without preserving structure.
        
        This method simply concatenates all paragraphs.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            Plain text without structure
        """
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return '\n\n'.join(paragraphs)
    
    def _extract_table(self, table) -> str:
        """
        Extract text from a table.
        
        Args:
            table: python-docx Table object
            
        Returns:
            Table as formatted text
        """
        rows = []
        
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                cells.append(cell_text if cell_text else '')
            rows.append(' | '.join(cells))
        
        if rows:
            # Add separator for markdown-style table
            separator = ' | '.join(['---'] * len(rows[0].split(' | ')))
            return '\n'.join([rows[0], separator] + rows[1:])
        
        return ''
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Text to clean
            
        Returns:
            Cleaned text
        """
        import re
        
        # Remove excessive whitespace
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    def get_metadata(self, filepath: Path | str) -> dict:
        """
        Extract metadata from a DOCX file.
        
        Args:
            filepath: Path to DOCX file
            
        Returns:
            Dictionary of metadata (author, title, created date, etc.)
        """
        path = Path(filepath)
        
        try:
            doc = Document(path)
            core_props = doc.core_properties
            
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'created': core_props.created.isoformat() if core_props.created else None,
                'modified': core_props.modified.isoformat() if core_props.modified else None,
                'last_modified_by': core_props.last_modified_by or '',
                'revision': core_props.revision or '',
                'version': core_props.version or '',
            }
            
            # Filter out empty values
            return {k: v for k, v in metadata.items() if v}
            
        except Exception as e:
            logger.warning("Failed to extract metadata from DOCX", filepath=str(path), error=str(e))
            return {}
